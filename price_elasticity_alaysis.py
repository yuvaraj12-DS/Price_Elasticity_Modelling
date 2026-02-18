from docx import Document
from docx.shared import Inches
import datetime
def generate_word_report(elasticity_df, cat_summary):
    doc = Document()
    doc.add_heading('Price Elasticity Modelling Report', 0)
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("This report summarizes the price elasticity analysis performed on the provided sales data, including data cleaning, exploratory analysis, elasticity estimation, and actionable business recommendations.")

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The analysis estimates price elasticity at the SKU and category level. Most SKUs are inelastic, but some categories and products are highly price-sensitive. Recommendations are provided for pricing strategy optimization."
    )

    doc.add_heading('2. SKU-Level Elasticity Summary', level=1)
    desc = elasticity_df['Elasticity'].describe().to_string()
    doc.add_paragraph('Elasticity statistics across all SKUs:')
    doc.add_paragraph(desc)

    doc.add_heading('3. Most and Least Price-Sensitive SKUs', level=1)
    doc.add_paragraph('Top 5 Most Price-Sensitive SKUs:')
    table1 = doc.add_table(rows=1, cols=2)
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Product Number'
    hdr_cells[1].text = 'Elasticity'
    for _, row in elasticity_df.nsmallest(5, 'Elasticity').iterrows():
        cells = table1.add_row().cells
        cells[0].text = str(row['Product Number'])
        cells[1].text = f"{row['Elasticity']:.2f}"
    doc.add_paragraph('Top 5 Least Price-Sensitive SKUs:')
    table2 = doc.add_table(rows=1, cols=2)
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Product Number'
    hdr_cells[1].text = 'Elasticity'
    for _, row in elasticity_df.nlargest(5, 'Elasticity').iterrows():
        cells = table2.add_row().cells
        cells[0].text = str(row['Product Number'])
        cells[1].text = f"{row['Elasticity']:.2f}"

    doc.add_heading('4. Category-Level Elasticity', level=1)
    doc.add_paragraph('Top 10 Categories by Average Elasticity:')
    table3 = doc.add_table(rows=1, cols=3)
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Avg Elasticity'
    hdr_cells[2].text = 'Recommendation'
    for _, row in cat_summary.head(10).iterrows():
        cells = table3.add_row().cells
        cells[0].text = str(row['L3 Category Name'])
        cells[1].text = f"{row['avg_elasticity']:.2f}"
        cells[2].text = row['Recommendation']

    doc.add_heading('5. Key Visualizations', level=1)
    doc.add_paragraph('The following plots summarize the main findings:')
    doc.add_picture('outputs/category_avg_elasticity.png', width=Inches(5.5))
    doc.add_paragraph('Figure 1: Average Price Elasticity by Category (Top 20)')
    doc.add_picture('outputs/top10_sensitive_skus.png', width=Inches(5.5))
    doc.add_paragraph('Figure 2: Top 10 Most Price-Sensitive SKUs')
    doc.add_picture('outputs/top10_insensitive_skus.png', width=Inches(5.5))
    doc.add_paragraph('Figure 3: Top 10 Least Price-Sensitive SKUs')

    doc.add_heading('6. Conclusions & Recommendations', level=1)
    doc.add_paragraph(
        "- Highly elastic SKUs and categories: Consider price reductions to boost sales.\n"
        "- Inelastic SKUs and categories: Price increases may be possible with minimal sales loss.\n"
        "- Outliers or positive elasticity: Review for data quality or special business cases.\n"
        "- Use these insights to inform pricing strategy and promotional planning."
    )

    doc.add_paragraph('End of Report')
    doc.save('outputs/Price_Elasticity_Report.docx')
    print('Word report saved to outputs/Price_Elasticity_Report.docx')
def generate_summary_report(elasticity_df, cat_summary, output_path="outputs/summary_report.txt"):
    with open(output_path, "w") as f:
        f.write("Price Elasticity Modelling Summary Report\n")
        f.write("="*40 + "\n\n")
        # Overall elasticity stats
        f.write("SKU-Level Elasticity Summary:\n")
        f.write(str(elasticity_df['Elasticity'].describe()) + "\n\n")
        # Most/least sensitive SKUs
        f.write("Most Price-Sensitive SKUs (most negative elasticity):\n")
        f.write(str(elasticity_df.nsmallest(5, 'Elasticity')[['Product Number', 'Elasticity']]) + "\n\n")
        f.write("Least Price-Sensitive SKUs (most positive elasticity):\n")
        f.write(str(elasticity_df.nlargest(5, 'Elasticity')[['Product Number', 'Elasticity']]) + "\n\n")
        # Category summary
        f.write("Category-Level Average Elasticity:\n")
        f.write(str(cat_summary[['L3 Category Name', 'avg_elasticity', 'Recommendation']].head(10)) + "\n\n")
        # Recommendations
        highly_elastic = cat_summary[cat_summary['avg_elasticity'] < -1]
        inelastic = cat_summary[(cat_summary['avg_elasticity'] >= -0.5) & (cat_summary['avg_elasticity'] < 0)]
        f.write("Recommendations:\n")
        if not highly_elastic.empty:
            f.write("- Highly elastic categories: Consider price cuts to boost sales.\n")
            f.write(str(highly_elastic[['L3 Category Name', 'avg_elasticity']]) + "\n\n")
        if not inelastic.empty:
            f.write("- Inelastic categories: Price changes have limited effect.\n")
            f.write(str(inelastic[['L3 Category Name', 'avg_elasticity']]) + "\n\n")
        f.write("- Review categories with positive or atypical elasticity for data quality or special business cases.\n")
        f.write("\nEnd of Report\n")
def visualize_results(elasticity_df, cat_summary):
    import matplotlib.pyplot as plt
    import seaborn as sns
    # Bar plot: average elasticity by category (top 20 for clarity)
    plt.figure(figsize=(12,8))
    sorted_cats = cat_summary.sort_values('avg_elasticity').head(20)
    ax = sns.barplot(x='avg_elasticity', y='L3 Category Name', data=sorted_cats, palette='coolwarm', edgecolor='black')
    plt.title('Average Price Elasticity by Category (Top 20)', fontsize=16, fontweight='bold')
    plt.xlabel('Average Price Elasticity', fontsize=13)
    plt.ylabel('Category', fontsize=13)
    plt.grid(axis='x', linestyle='--', alpha=0.6)
    # Add value labels
    for p in ax.patches:
        ax.annotate(f"{p.get_width():.2f}", (p.get_width(), p.get_y() + p.get_height()/2),
                    ha='left', va='center', fontsize=11, color='black', xytext=(5,0), textcoords='offset points')
    plt.tight_layout()
    plt.savefig('outputs/category_avg_elasticity.png', dpi=300, bbox_inches='tight')
    plt.show()

    # Bar plot: top 10 most/least elastic SKUs (vertical, categorical x-axis)
    top10 = elasticity_df.nsmallest(10, 'Elasticity').copy()
    bottom10 = elasticity_df.nlargest(10, 'Elasticity').copy()
    top10['Product Number'] = top10['Product Number'].astype(str)
    bottom10['Product Number'] = bottom10['Product Number'].astype(str)

    # Most price-sensitive SKUs
    plt.figure(figsize=(12,6))
    ax = sns.barplot(x='Product Number', y='Elasticity', data=top10, color='royalblue', edgecolor='black')
    plt.title('Top 10 Most Price-Sensitive SKUs', fontsize=15, fontweight='bold')
    plt.xlabel('Product Number', fontsize=12)
    plt.ylabel('Elasticity', fontsize=12)
    plt.grid(axis='y', linestyle='--', alpha=0.6)
    # Add value labels above bars
    for p in ax.patches:
        ax.annotate(f"{p.get_height():.2f}", (p.get_x() + p.get_width()/2, p.get_height()),
                    ha='center', va='bottom', fontsize=10, color='black', xytext=(0,3), textcoords='offset points')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('outputs/top10_sensitive_skus.png', dpi=300, bbox_inches='tight')
    plt.show()

    # Least price-sensitive SKUs
    plt.figure(figsize=(12,6))
    ax = sns.barplot(x='Product Number', y='Elasticity', data=bottom10, color='tomato', edgecolor='black')
    plt.title('Top 10 Least Price-Sensitive SKUs', fontsize=15, fontweight='bold')
    plt.xlabel('Product Number', fontsize=12)
    plt.ylabel('Elasticity', fontsize=12)
    plt.grid(axis='y', linestyle='--', alpha=0.6)
    for p in ax.patches:
        ax.annotate(f"{p.get_height():.2f}", (p.get_x() + p.get_width()/2, p.get_height()),
                    ha='center', va='bottom', fontsize=10, color='black', xytext=(0,3), textcoords='offset points')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('outputs/top10_insensitive_skus.png', dpi=300, bbox_inches='tight')
    plt.show()

# --- Price Elasticity Modelling Script ---
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.linear_model import LinearRegression
from scipy.stats import zscore

def load_and_clean_data(filepath, sheet_name, skiprows=5):
    data = pd.read_excel(filepath, sheet_name=sheet_name, skiprows=skiprows)
    # Data quality checks and cleaning
    print(data.head(10))
    print(data.columns)
    print(data.info())
    print(data.describe())
    print("Missing values:\n", data.isnull().sum())
    print("Is sorted by date?", data['Start of Week'].is_monotonic_increasing)
    data['price_zscore'] = zscore(data['Current (3.2.26) national price excl GST'])
    print("Extreme price outliers:\n", data[data['price_zscore'].abs() > 3])
    print("Duplicate rows:", data.duplicated().sum())
    gm_missing = data[data['GM %'].isnull()]
    print("Missing GM % rows:\n", gm_missing[['Product Number','Product Name','Start of Week','Sales $','GM $']].head(10))
    asp_missing = data[data['Average Sell Price'].isnull()]
    print("Missing Average Sell Price rows:\n", asp_missing[['Product Number','Product Name','Start of Week','Sales $','Sales Qty']].head(10))
    data['GM %'] = data['GM %'].fillna(np.nan)
    data['Average Sell Price'] = data['Average Sell Price'].fillna(np.nan)
    data['No Sales Flag'] = (data['Sales $'] == 0) | (data['Sales Qty'] == 0)
    mask_gm = data['GM %'].isnull() & (data['Sales $'] != 0)
    data.loc[mask_gm, 'GM %'] = (data.loc[mask_gm, 'GM $'] / data.loc[mask_gm, 'Sales $']) * 100
    mask_asp = data['Average Sell Price'].isnull() & (data['Sales Qty'] > 0)
    data.loc[mask_asp, 'Average Sell Price'] = data.loc[mask_asp, 'Sales $'] / data.loc[mask_asp, 'Sales Qty']
    data['L3 Category Name'] = data['L3 Category Name'].fillna("Unknown")
    data = data.sort_values('Start of Week').reset_index(drop=True)
    q_low = data['Current (3.2.26) national price excl GST'].quantile(0.01)
    q_high = data['Current (3.2.26) national price excl GST'].quantile(0.99)
    data['Current (3.2.26) national price excl GST'] = np.clip(
        data['Current (3.2.26) national price excl GST'], q_low, q_high)
    return data

def exploratory_data_analysis(data):
    # 1. Price and Sales Trends Over Time
    plt.figure(figsize=(12, 5))
    plt.plot(data['Start of Week'], data['Current (3.2.26) national price excl GST'], label='Price')
    plt.ylabel('Price')
    plt.xlabel('Date')
    plt.title('Price Trend Over Time')
    plt.legend()
    plt.tight_layout()
    plt.show()

    plt.figure(figsize=(12, 5))
    plt.plot(data['Start of Week'], data['Sales $'], label='Sales $', color='orange')
    plt.ylabel('Sales $')
    plt.xlabel('Date')
    plt.title('Sales Trend Over Time')
    plt.legend()
    plt.tight_layout()
    plt.show()

    # 2. Price vs. Quantity Sold
    plt.figure(figsize=(8, 6))
    sns.scatterplot(x='Current (3.2.26) national price excl GST', y='Sales Qty', data=data)
    plt.xlabel('Price')
    plt.ylabel('Quantity Sold')
    plt.title('Price vs. Quantity Sold')
    plt.tight_layout()
    plt.show()

    # 3. Correlation
    correlation = data[['Current (3.2.26) national price excl GST', 'Sales Qty']].corr()
    print('Correlation between price and quantity sold:')
    print(correlation)

    # SKU-specific EDA example
    sku_sample = data[data['Product Number'] == 1000131]  # example SKU
    plt.figure(figsize=(12,6))
    plt.plot(sku_sample['Start of Week'], sku_sample['Sales Qty'], label='Sales Qty')
    plt.plot(sku_sample['Start of Week'], sku_sample['Average Sell Price'], label='Avg Sell Price')
    plt.legend()
    plt.title("Sales Qty & Avg Price over Time (SKU 1000131)")
    plt.show()

def estimate_price_elasticity(data):
    elasticity_results = []
    sku_list = data['Product Number'].unique()
    for sku in sku_list:
        sku_data = data[(data['Product Number'] == sku) &
                       (data['Sales Qty'] > 0) &
                       (data['Current (3.2.26) national price excl GST'] > 0)]
        if len(sku_data) < 10:
            continue  # skip SKUs with too little data
        sku_data = sku_data.copy()
        sku_data['log_qty'] = np.log(sku_data['Sales Qty'])
        sku_data['log_price'] = np.log(sku_data['Current (3.2.26) national price excl GST'])
        X = sku_data[['log_price']]
        y = sku_data['log_qty']
        model = LinearRegression().fit(X, y)
        elasticity = model.coef_[0]
        elasticity_results.append({'Product Number': sku,
                                   'Elasticity': elasticity,
                                   'n_obs': len(sku_data)})
    elasticity_df = pd.DataFrame(elasticity_results)
    print('Sample of SKU-level price elasticity estimates:')
    print(elasticity_df.head(10))
    print('Elasticity summary:')
    print(elasticity_df['Elasticity'].describe())
    plt.figure(figsize=(8,6))
    sns.histplot(elasticity_df['Elasticity'], bins=30, kde=True)
    plt.title('Distribution of SKU-level Price Elasticities')
    plt.xlabel('Price Elasticity')
    plt.ylabel('Number of SKUs')
    plt.tight_layout()
    plt.show()
    return elasticity_df

def segment_analysis_by_category(elasticity_df, data):
    # Merge elasticity with category info
    merged = elasticity_df.merge(data[['Product Number', 'L3 Category Name']].drop_duplicates(), on='Product Number', how='left')
    # Group by category
    cat_summary = merged.groupby('L3 Category Name').agg(
        avg_elasticity=('Elasticity', 'mean'),
        count=('Product Number', 'count')
    ).reset_index().sort_values('avg_elasticity')
    print('\nAverage price elasticity by category:')
    print(cat_summary)
    # Add recommendation
    def cat_recommendation(e):
        if e < -1:
            return 'Highly elastic: Consider price cuts.'
        elif e < -0.5:
            return 'Moderately elastic: Price changes will impact sales.'
        elif e < 0:
            return 'Inelastic: Price changes have limited effect.'
        else:
            return 'Atypical/positive elasticity: Review category.'
    cat_summary['Recommendation'] = cat_summary['avg_elasticity'].apply(cat_recommendation)
    print('\nCategory-level recommendations:')
    print(cat_summary[['L3 Category Name', 'avg_elasticity', 'Recommendation']])
    return cat_summary
    # Merge elasticity with category info
    merged = elasticity_df.merge(data[['Product Number', 'L3 Category Name']].drop_duplicates(), on='Product Number', how='left')
    # Group by category
    cat_summary = merged.groupby('L3 Category Name').agg(
        avg_elasticity=('Elasticity', 'mean'),
        count=('Product Number', 'count')
    ).reset_index().sort_values('avg_elasticity')
    print('\nAverage price elasticity by category:')
    print(cat_summary)
    # Add recommendation
    def cat_recommendation(e):
        if e < -1:
            return 'Highly elastic: Consider price cuts.'
        elif e < -0.5:
            return 'Moderately elastic: Price changes will impact sales.'
        elif e < 0:
            return 'Inelastic: Price changes have limited effect.'
        else:
            return 'Atypical/positive elasticity: Review category.'
    cat_summary['Recommendation'] = cat_summary['avg_elasticity'].apply(cat_recommendation)
    print('\nCategory-level recommendations:')
    print(cat_summary[['L3 Category Name', 'avg_elasticity', 'Recommendation']])

def actionable_insights(elasticity_df):
    # 1. Most and least price-sensitive SKUs
    top_sensitive = elasticity_df.sort_values('Elasticity').head(10)
    top_insensitive = elasticity_df.sort_values('Elasticity', ascending=False).head(10)
    print('\nMost price-sensitive SKUs (most negative elasticity):')
    print(top_sensitive)
    print('\nLeast price-sensitive SKUs (most positive elasticity):')
    print(top_insensitive)
    # 2. Flag outliers (positive elasticity or extreme values)
    outliers = elasticity_df[(elasticity_df['Elasticity'] > 0.5) | (elasticity_df['Elasticity'] < -2)]
    print('\nElasticity outliers (possible data issues or special cases):')
    print(outliers)
    # 3. Business recommendation summary
    def elasticity_recommendation(e):
        if e < -1:
            return 'Highly elastic: Consider price cuts to boost sales.'
        elif e < -0.5:
            return 'Moderately elastic: Price changes will impact sales.'
        elif e < 0:
            return 'Inelastic: Price changes have limited effect.'
        else:
            return 'Atypical/positive elasticity: Review data or product.'
    elasticity_df['Recommendation'] = elasticity_df['Elasticity'].apply(elasticity_recommendation)
    print('\nSample business recommendations:')
    print(elasticity_df[['Product Number', 'Elasticity', 'Recommendation']].head(10))


import os
if __name__ == "__main__":
    # 1. Load and clean data
    data = load_and_clean_data(
        'Copy of 3.02.26 Monthly Trader March 26 Skus Last 365 days_with PI inputs.xlsx',
        sheet_name="1. Monthly Trader last 365 days",
        skiprows=5
    )
    # 2. Exploratory Data Analysis
    exploratory_data_analysis(data)
    # 3. Estimate price elasticity
    elasticity_df = estimate_price_elasticity(data)
    # 4. Actionable insights
    actionable_insights(elasticity_df)
    # 5. Segment analysis by category
    cat_summary = segment_analysis_by_category(elasticity_df, data)

    # 6. Export results to Excel
    os.makedirs('outputs', exist_ok=True)
    print("\nExporting results to Excel files...")
    elasticity_df.to_excel("outputs/sku_level_elasticity.xlsx", index=False)
    cat_summary.to_excel("outputs/category_level_elasticity.xlsx", index=False)
    print("Export complete: 'outputs/sku_level_elasticity.xlsx' and 'outputs/category_level_elasticity.xlsx'")

    # 7. Visualize results
    visualize_results(elasticity_df, cat_summary)
    # 8. Generate summary report
    generate_summary_report(elasticity_df, cat_summary)
    print("Summary report saved to outputs/summary_report.txt")
    # 9. Generate Word report
    generate_word_report(elasticity_df, cat_summary)

# 2. Price vs. Quantity Sold
plt.figure(figsize=(8, 6))
sns.scatterplot(x='Current (3.2.26) national price excl GST', y='Sales Qty', data=data)
plt.xlabel('Price')
plt.ylabel('Quantity Sold')
plt.title('Price vs. Quantity Sold')
plt.tight_layout()
plt.show()

# 3. Correlation
correlation = data[['Current (3.2.26) national price excl GST', 'Sales Qty']].corr()
print('Correlation between price and quantity sold:')
print(correlation)

print(data.isnull().sum())
print(data['Start of Week'].is_monotonic_increasing)
print(data.describe())


# --- Price Elasticity Estimation (SKU Level) ---
from sklearn.linear_model import LinearRegression

elasticity_results = []
sku_list = data['Product Number'].unique()

for sku in sku_list:
    sku_data = data[(data['Product Number'] == sku) &
                   (data['Sales Qty'] > 0) &
                   (data['Current (3.2.26) national price excl GST'] > 0)]
    if len(sku_data) < 10:
        continue  # skip SKUs with too little data
    # Log transform
    sku_data = sku_data.copy()
    sku_data['log_qty'] = np.log(sku_data['Sales Qty'])
    sku_data['log_price'] = np.log(sku_data['Current (3.2.26) national price excl GST'])
    X = sku_data[['log_price']]
    y = sku_data['log_qty']
    model = LinearRegression().fit(X, y)
    elasticity = model.coef_[0]
    elasticity_results.append({'Product Number': sku,
                               'Elasticity': elasticity,
                               'n_obs': len(sku_data)})

elasticity_df = pd.DataFrame(elasticity_results)
print('Sample of SKU-level price elasticity estimates:')
print(elasticity_df.head(10))
print('Elasticity summary:')
print(elasticity_df['Elasticity'].describe())

# Visualize distribution of elasticities
plt.figure(figsize=(8,6))
sns.histplot(elasticity_df['Elasticity'], bins=30, kde=True)
plt.title('Distribution of SKU-level Price Elasticities')
plt.xlabel('Price Elasticity')
plt.ylabel('Number of SKUs')
plt.tight_layout()
plt.show()
